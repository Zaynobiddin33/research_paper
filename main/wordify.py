from datetime import datetime
from docx import Document
from docxcompose.composer import Composer

def add_template(template_path, existing_path, output_path):
    template = Document(template_path)
    composer = Composer(template)
    existing_doc = Document(existing_path)

    composer.append(existing_doc)
    composer.save(output_path)
    print(f"✅ Template added (cross-platform version): {output_path}")
    return output_path


def fill_template(template_path, replacements, output_path):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(f"{{{{{key}}}}}", value)

    # Replace placeholders inside tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", value)

    doc.save(output_path)
    print(f"✅ Filled template saved as: {output_path}")
    return output_path


